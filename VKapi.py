import urllib.request
import os
import vk
from time import sleep
from progress.bar import IncrementalBar
import datetime
import docx
from argparse import ArgumentParser
import shutil
from dateutil.relativedelta import relativedelta

parser = ArgumentParser(description='A nice tool for creating friends album')
required_group = parser.add_argument_group('required arguments')
required_group.add_argument('-id', '--user_id',
                            help='user id (required)')
parser.add_argument('-d', '--include_deactivated', action='store_true',
                    help='include deactivated users')
parser.add_argument('-name', '--album_filename', action='store_true', default='album',
                    help='specify album filename')

args = parser.parse_args()

# https://oauth.vk.com/authorize?client_id=7815028&display=page&redirect_uri=http://oauth.vk.com/blank.html&scope=friends&response_type=token&v=5.130


access_token = '45ce383bbb25fa1de73091c43cbf47a6f415140e8ea1536ffc966e1e63e4dc25e34d18612ec514d6ddcf6'
session = vk.Session(access_token=access_token)
vkapi = vk.API(session)
ids = vkapi.friends.get(user_id=args.user_id, v='5.130')

bar = IncrementalBar('Processing... ', max = ids['count'])


counter = 0
os.mkdir(os.getcwd()+'/tmp')
document = docx.Document()
document.save(os.getcwd()+'/' + args.album_filename + '.docx')
for id in ids['items']:
    info = vkapi.users.get(user_id=id, fields='photo_50, sex, bdate', name_case='nom', v='5.130')
    bar.next()
    counter+=1
    if counter == 6:
        sleep(5)
        counter = 0
    if 'deactivated' in info[0] and not args.include_deactivated:
        continue

    urllib.request.urlretrieve(info[0]['photo_50'], os.getcwd()+'/tmp/'+str(info[0]['id'])+'.jpg')
    document.add_picture(os.getcwd()+'/tmp/'+str(info[0]['id'])+'.jpg')
    document.add_paragraph(info[0]['first_name'] + ' ' + info[0]['last_name'])
    try:
        document.add_paragraph('Дата рождения: '+info[0]['bdate'])
        bdata = datetime.date(int(info[0]['bdate'].split('.')[2]), int(info[0]['bdate'].split('.')[1]), int(info[0]['bdate'].split('.')[0]))
        rdelta = relativedelta(datetime.datetime.now(), bdata)
        document.add_paragraph('Возраст: ' + str(rdelta.years))
    except KeyError:
        document.add_paragraph('Дата рождения: Не указана')
    except IndexError:
        document.add_paragraph('Возраст: Год рождения не указан')
    except ValueError:
        print(datetime.datetime.now())
        print(info[0]['bdate'])
    document.add_paragraph('Ж' if info[0]['sex'] == 1 else 'М')
    document.save(os.getcwd()+'/album.docx')

shutil.rmtree(os.getcwd()+'/tmp')
bar.finish()